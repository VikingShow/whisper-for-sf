from faster_whisper import WhisperModel
from opencc import OpenCC
from docx import Document
from docx.shared import Pt
import re
import os
import time

# ========== 配置 ==========
AUDIO_FILE = "星雨街.m4a"
MODEL_SIZE = "large-v3"    # base / small / medium / large-v3
DEVICE = "cuda"
COMPUTE_TYPE = "float16"
SEGMENT_LENGTH = 300        # 每段约 300 秒 = 5 分钟，可根据需要调整
# =========================

def format_time(seconds):
    s = int(seconds)
    h = s // 3600
    m = (s % 3600) // 60
    sec = s % 60
    return f"{h:02d}:{m:02d}:{sec:02d}"

def add_basic_punctuation(text):
    """
    简单中文标点优化：
    - 去掉多余空格
    - 在没有句号或问号的地方添加句号
    """
    text = re.sub(r"\s+", " ", text)
    text = text.strip()
    # 遍历文本，如果以中文字符结尾且没有标点，补充句号
    if text and text[-1] not in "。！？":
        text += "。"
    return text

def main():
    start_time = time.time()

    # 自动生成 Word 文件名
    base_name = os.path.splitext(os.path.basename(AUDIO_FILE))[0]
    WORD_OUTPUT = f"{base_name}.docx"

    cc = OpenCC("t2s")
    model = WhisperModel(MODEL_SIZE, device=DEVICE, compute_type=COMPUTE_TYPE)

    print(">>> 开始加载模型并转写音频...")
    segments, info = model.transcribe(
        AUDIO_FILE,
        beam_size=5,
        language="zh"
    )
    print(f">>> 音频时长：{info.duration:.2f} 秒")

    # 合并分段
    merged_segments = []
    temp_text = ""
    temp_start = None
    temp_end = None

    for seg in segments:
        text = cc.convert(seg.text.strip())
        text = add_basic_punctuation(text)
        if temp_start is None:
            temp_start = seg.start
        temp_end = seg.end
        temp_text += text + " "
        if (temp_end - temp_start) >= SEGMENT_LENGTH:
            merged_segments.append((temp_start, temp_end, temp_text.strip()))
            temp_text = ""
            temp_start = None
            temp_end = None

    if temp_text:
        merged_segments.append((temp_start, temp_end, temp_text.strip()))

    # 写入 Word
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style.font.size = Pt(12)
    doc.add_heading(base_name, level=1)  # 标题为音频名

    for i, (start, end, text) in enumerate(merged_segments):
        doc.add_heading(f"第 {i+1} 段（{format_time(start)} → {format_time(end)}）", level=2)
        doc.add_paragraph(text)

    doc.save(WORD_OUTPUT)
    end_time = time.time()
    duration = end_time - start_time

    print(f"✅ 已写入 Word：{WORD_OUTPUT}")
    print(f"⏱ 转写耗时：{duration:.2f} 秒（约 {duration/60:.2f} 分钟）")

if __name__ == "__main__":
    main()
