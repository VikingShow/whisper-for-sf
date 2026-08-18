"""
Unit tests for formatters.py.
"""
import os
import tempfile

import pytest

pytest.importorskip("docx")
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from formatters import DocxFormatter, FormatterFactory, SrtFormatter, TxtFormatter


@pytest.fixture
def sample_segments():
    return [
        (0.0, 10.5, "这是第一段文本。"),
        (10.5, 25.3, "这是第二段文本，内容更长一些。"),
        (25.3, 40.0, "这是第三段文本，包含更多信息。"),
    ]


class TestFormatterFactory:
    def test_create_docx_formatter(self):
        formatter = FormatterFactory.create("docx", "Test")
        assert isinstance(formatter, DocxFormatter)

    def test_create_txt_formatter(self):
        formatter = FormatterFactory.create("txt", "Test")
        assert isinstance(formatter, TxtFormatter)

    def test_create_srt_formatter(self):
        formatter = FormatterFactory.create("srt", "Test")
        assert isinstance(formatter, SrtFormatter)

    def test_create_invalid_formatter(self):
        with pytest.raises(ValueError, match="不支持的格式"):
            FormatterFactory.create("invalid", "Test")

    def test_get_supported_formats(self):
        formats = FormatterFactory.get_supported_formats()
        assert "docx" in formats
        assert "txt" in formats
        assert "srt" in formats


class TestTxtFormatter:
    def test_txt_format(self, sample_segments):
        formatter = TxtFormatter("测试标题")

        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            output_path = f.name

        try:
            formatter.format(sample_segments, output_path)
            assert os.path.exists(output_path)
            with open(output_path, "r", encoding="utf-8") as f:
                content = f.read()
            assert "测试标题" in content
            assert "第一段文本" in content
            assert "00:00:00" in content
        finally:
            if os.path.exists(output_path):
                os.remove(output_path)


class TestSrtFormatter:
    def test_srt_format(self, sample_segments):
        formatter = SrtFormatter("测试标题")

        with tempfile.NamedTemporaryFile(mode="w", suffix=".srt", delete=False, encoding="utf-8") as f:
            output_path = f.name

        try:
            formatter.format(sample_segments, output_path)
            assert os.path.exists(output_path)
            with open(output_path, "r", encoding="utf-8") as f:
                content = f.read()
            assert "1\n" in content
            assert "-->" in content
            assert "第一段文本" in content
        finally:
            if os.path.exists(output_path):
                os.remove(output_path)


class TestDocxFormatter:
    def test_docx_format(self, sample_segments):
        formatter = DocxFormatter("测试标题")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name

        try:
            formatter.format(sample_segments, output_path)
            assert os.path.exists(output_path)
            assert os.path.getsize(output_path) > 0
        finally:
            if os.path.exists(output_path):
                os.remove(output_path)

    def test_docx_detects_verse_block(self):
        segments = [
            (0.0, 2.0, "首先我们来解释这个祈祷文。"),
            (2.0, 3.0, "遍满虚空尽边际，上师本尊空行众。"),
            (3.0, 4.0, "诸佛正法圣僧前，我与六道敬皈依。"),
            (4.0, 6.0, "以上两句就是颂词的核心。"),
        ]
        formatter = DocxFormatter("讲记")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name

        try:
            formatter.format(segments, output_path)
            doc = Document(output_path)
            verse_para = None
            for para in doc.paragraphs:
                if "遍满虚空尽边际" in para.text:
                    verse_para = para
                    break
            assert verse_para is not None
            assert verse_para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
        finally:
            if os.path.exists(output_path):
                os.remove(output_path)

    def test_docx_applies_target_margins(self):
        formatter = DocxFormatter("测试标题")
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        try:
            formatter.format([(0.0, 1.0, "正文。")], output_path)
            doc = Document(output_path)
            section = doc.sections[0]
            assert round(section.top_margin.cm, 2) == 2.54
            assert round(section.bottom_margin.cm, 2) == 2.54
            assert round(section.left_margin.cm, 2) == 3.18
            assert round(section.right_margin.cm, 2) == 3.18
        finally:
            if os.path.exists(output_path):
                os.remove(output_path)

    def test_docx_uses_template_and_clears_template_body(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            template_path = f.name
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name

        try:
            template_doc = Document()
            template_doc.add_paragraph("模板原有内容")
            template_doc.save(template_path)

            formatter = DocxFormatter("新导出标题", options={"docx_template": template_path})
            formatter.format([(0.0, 1.0, "正文内容。")], output_path)

            doc = Document(output_path)
            texts = [p.text for p in doc.paragraphs if p.text.strip()]
            assert "模板原有内容" not in texts
            assert any("新导出标题" in t for t in texts)
        finally:
            if os.path.exists(template_path):
                os.remove(template_path)
            if os.path.exists(output_path):
                os.remove(output_path)
