"""
输出格式化模块
支持多种输出格式：Word、纯文本、字幕（SRT）
"""
from __future__ import annotations

import logging
import os
import re
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt

from utils import format_time, format_timestamp

logger = logging.getLogger(__name__)
Segment = Tuple[float, float, str]


class BaseFormatter:
    """基础格式化器"""
    
    def __init__(self, title: str, options: Optional[Dict[str, object]] = None):
        self.title = title
        self.options = options or {}
    
    def format(self, segments: List[Segment], output_path: str) -> None:
        """
        格式化并保存输出
        
        Args:
            segments: 片段列表 [(start_time, end_time, text), ...]
            output_path: 输出文件路径
        """
        raise NotImplementedError


class DocxFormatter(BaseFormatter):
    """Word 文档格式化器"""

    _VERSE_HINT_WORDS = ("偈", "偈颂", "颂词", "祈祷文", "念诵", "仪轨")
    
    def format(self, segments: List[Segment], output_path: str) -> None:
        """
        生成 Word 文档
        
        Args:
            segments: 片段列表 [(start_time, end_time, text), ...]
            output_path: 输出文件路径
        """
        try:
            doc = self._build_document()
            self._apply_page_margins(doc)
            prose_style = self._resolve_style(
                doc,
                str(self.options.get("docx_prose_style", "")).strip(),
                fallback_candidates=["Normal", "Body Text"],
            )
            verse_style = self._resolve_style(
                doc,
                str(self.options.get("docx_verse_style", "")).strip(),
                fallback_candidates=["Body Text", "Intense Quote", "Quote", prose_style],
            )

            title_para = doc.add_paragraph(self.title, style=prose_style)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in title_para.runs:
                run.bold = True
                run.font.size = Pt(16)

            total_duration = segments[-1][1] if segments else 0
            meta = doc.add_paragraph()
            meta.add_run(f"总时长: {format_time(total_duration)}\n")
            meta.add_run(f"转写片段: {len(segments)}")
            meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            show_timestamps = bool(self.options.get("docx_add_timestamps", True))
            ts_interval = int(self.options.get("docx_timestamp_interval", 300))
            next_ts_time = 0.0  # 下一个时间戳触发时间点

            blocks = self._build_blocks(segments)
            for block in blocks:
                block_type = block[0]
                start = block[1]
                end = block[2]

                # 每隔 ts_interval 秒插入独立时间戳段落
                if show_timestamps and start >= next_ts_time:
                    ts_para = doc.add_paragraph(
                        f"【{format_time(start)}】", style=prose_style
                    )
                    ts_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in ts_para.runs:
                        run.bold = True
                        run.font.size = Pt(10)
                    ts_para.paragraph_format.space_before = Pt(12)
                    ts_para.paragraph_format.space_after = Pt(4)
                    next_ts_time = (start // ts_interval + 1) * ts_interval

                if block_type == "verse":
                    lines: List[str] = block[3]  # type: ignore[assignment]
                    for line in lines:
                        para = doc.add_paragraph(line, style=verse_style)
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        para.paragraph_format.line_spacing = 1.2
                        para.paragraph_format.space_after = Pt(0)
                    doc.add_paragraph("")
                    continue

                text: str = block[3]  # type: ignore[assignment]
                para = doc.add_paragraph(text, style=prose_style)
                para.paragraph_format.line_spacing = 1.5
                para.paragraph_format.first_line_indent = Pt(24)
                para.paragraph_format.space_after = Pt(8)

            # 保存文档
            doc.save(output_path)
            logger.info(f"已保存 Word 文档: {output_path}")
            
        except Exception as e:
            logger.error(f"生成 Word 文档失败: {e}")
            raise

    def _apply_page_margins(self, doc: Document) -> None:
        top_cm = float(self.options.get("docx_margin_top_cm", 2.54))
        bottom_cm = float(self.options.get("docx_margin_bottom_cm", 2.54))
        left_cm = float(self.options.get("docx_margin_left_cm", 3.18))
        right_cm = float(self.options.get("docx_margin_right_cm", 3.18))

        for section in doc.sections:
            section.top_margin = Cm(top_cm)
            section.bottom_margin = Cm(bottom_cm)
            section.left_margin = Cm(left_cm)
            section.right_margin = Cm(right_cm)

    @staticmethod
    def _normalize_text(text: str) -> str:
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = text.replace("\u3000", " ").strip()
        text = re.sub(r"[ \t]+", " ", text)
        return text

    def _build_document(self) -> Document:
        template_path = str(self.options.get("docx_template", "")).strip()
        if template_path:
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"DOCX 模板不存在: {template_path}")
            doc = Document(template_path)
            self._clear_document_content(doc)
            return doc

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "SimSun"
        style.font.size = Pt(12)
        return doc

    @staticmethod
    def _clear_document_content(doc: Document) -> None:
        # 仅删除正文节点，保留页面设置和样式定义（模板化核心）
        body = doc._element.body  # type: ignore[attr-defined]
        for element in list(body):
            if element.tag.endswith("sectPr"):
                continue
            body.remove(element)

    @staticmethod
    def _resolve_style(doc: Document, preferred: str, fallback_candidates: List[str]) -> str:
        style_names = {style.name for style in doc.styles}
        if preferred and preferred in style_names:
            return preferred
        for name in fallback_candidates:
            if name in style_names:
                return name
        return "Normal"

    @classmethod
    def _is_verse_intro_line(cls, text: str) -> bool:
        stripped = text.strip("：:，。！？!?, ")
        return len(stripped) <= 2

    @classmethod
    def _is_probably_verse_line(cls, text: str) -> bool:
        candidate = text.strip()
        if not candidate:
            return False
        if candidate.endswith(("：", ":")):
            return False
        if len(candidate) > 36:
            return False
        if any(ch in candidate for ch in ("？", "?", "；")) and len(candidate) > 20:
            return False

        has_break_punc = any(ch in candidate for ch in ("，", "。", "、"))
        short_plain_line = len(candidate) <= 16
        has_alpha_or_digit = bool(re.search(r"[A-Za-z0-9]", candidate))
        if has_alpha_or_digit and not has_break_punc:
            return False
        return has_break_punc or short_plain_line

    @classmethod
    def _split_prose_paragraphs(cls, text: str, max_len: int = 140) -> List[str]:
        normalized = cls._normalize_text(text)
        if not normalized:
            return []

        sentences = re.split(r"(?<=[。！？!?])", normalized)
        sentences = [sentence.strip() for sentence in sentences if sentence.strip()]
        if not sentences:
            return [normalized]

        paragraphs: List[str] = []
        current = ""
        for sentence in sentences:
            if not current:
                current = sentence
                continue
            if len(current) + len(sentence) <= max_len:
                current += sentence
                continue
            paragraphs.append(current)
            current = sentence
        if current:
            paragraphs.append(current)
        return paragraphs

    @classmethod
    def _build_blocks(cls, segments: List[Segment]) -> List[Tuple[str, float, float, object]]:
        entries: List[Segment] = []
        for start, end, text in segments:
            clean = cls._normalize_text(text)
            if clean:
                entries.append((start, end, clean))

        blocks: List[Tuple[str, float, float, object]] = []
        i = 0
        hint_active = False
        while i < len(entries):
            start, end, text = entries[i]
            if any(word in text for word in cls._VERSE_HINT_WORDS):
                hint_active = True

            next_text = entries[i + 1][2] if i + 1 < len(entries) else ""
            can_start_verse = (
                cls._is_probably_verse_line(text) and cls._is_probably_verse_line(next_text)
            ) or (
                cls._is_verse_intro_line(text)
                and cls._is_probably_verse_line(next_text)
                and i + 2 < len(entries)
                and cls._is_probably_verse_line(entries[i + 2][2])
            )

            if can_start_verse or (hint_active and cls._is_probably_verse_line(text)):
                verse_lines: List[str] = []
                verse_start = start
                j = i
                if cls._is_verse_intro_line(entries[j][2]) and j + 1 < len(entries) and cls._is_probably_verse_line(entries[j + 1][2]):
                    verse_lines.append(entries[j][2])
                    j += 1

                while j < len(entries) and cls._is_probably_verse_line(entries[j][2]):
                    verse_lines.append(entries[j][2])
                    j += 1

                verse_line_count = sum(1 for line in verse_lines if cls._is_probably_verse_line(line))
                if verse_line_count >= 2:
                    verse_end = entries[j - 1][1]
                    blocks.append(("verse", verse_start, verse_end, verse_lines))
                    i = j
                    hint_active = False
                    continue

            for paragraph in cls._split_prose_paragraphs(text):
                blocks.append(("prose", start, end, paragraph))
            i += 1

        return blocks


class TxtFormatter(BaseFormatter):
    """纯文本格式化器"""
    
    def format(self, segments: List[Segment], output_path: str) -> None:
        """
        生成纯文本文件
        
        Args:
            segments: 片段列表 [(start_time, end_time, text), ...]
            output_path: 输出文件路径
        """
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                # 写入标题和元数据
                f.write(f"{'=' * 60}\n")
                f.write(f"{self.title}\n")
                f.write(f"{'=' * 60}\n\n")
                
                total_duration = segments[-1][1] if segments else 0
                f.write(f"总时长: {format_time(total_duration)}\n")
                f.write(f"段落数: {len(segments)}\n")
                f.write(f"{'-' * 60}\n\n")
                
                # 写入各个片段
                for i, (start, end, text) in enumerate(segments, 1):
                    f.write(f"【第 {i} 段】{format_time(start)} → {format_time(end)}\n")
                    f.write(f"{text}\n\n")
                    f.write(f"{'-' * 60}\n\n")
            
            logger.info(f"已保存纯文本文件: {output_path}")
            
        except Exception as e:
            logger.error(f"生成纯文本文件失败: {e}")
            raise


class SrtFormatter(BaseFormatter):
    """SRT 字幕格式化器"""
    
    def format(self, segments: List[Segment], output_path: str) -> None:
        """
        生成 SRT 字幕文件
        
        Args:
            segments: 片段列表 [(start_time, end_time, text), ...]
            output_path: 输出文件路径
        """
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                for i, (start, end, text) in enumerate(segments, 1):
                    # SRT 格式：
                    # 序号
                    # 开始时间 --> 结束时间
                    # 字幕文本
                    # 空行
                    f.write(f"{i}\n")
                    f.write(f"{format_timestamp(start)} --> {format_timestamp(end)}\n")
                    f.write(f"{text}\n\n")
            
            logger.info(f"已保存 SRT 字幕文件: {output_path}")
            
        except Exception as e:
            logger.error(f"生成 SRT 字幕文件失败: {e}")
            raise


class MarkdownFormatter(BaseFormatter):
    """Markdown 文本格式化器"""
    
    def format(self, segments: List[Segment], output_path: str) -> None:
        """
        生成 Markdown 文件
        
        Args:
            segments: 片段列表 [(start_time, end_time, text), ...]
            output_path: 输出文件路径
        """
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                # 标题
                f.write(f"# {self.title}\n\n")
                
                # 元数据
                total_duration = segments[-1][1] if segments else 0
                f.write(f"> 总时长: {format_time(total_duration)}  ")
                f.write(f"段落数: {len(segments)}\n\n")
                f.write("---\n\n")
                
                # 内容
                for i, (start, end, text) in enumerate(segments, 1):
                    f.write(f"## 第 {i} 段 ({format_time(start)} → {format_time(end)})\n\n")
                    f.write(f"{text}\n\n")
            
            logger.info(f"已保存 Markdown 文件: {output_path}")
        except Exception as e:
            logger.error(f"生成 Markdown 文件失败: {e}")
            raise


class FormatterFactory:
    """格式化器工厂"""
    
    _formatters = {
        'docx': DocxFormatter,
        'txt': TxtFormatter,
        'srt': SrtFormatter,
        'md': MarkdownFormatter,
        'markdown': MarkdownFormatter,
    }
    
    @classmethod
    def create(
        cls,
        format_type: str,
        title: str,
        options: Optional[Dict[str, object]] = None,
    ) -> BaseFormatter:
        """
        创建格式化器
        
        Args:
            format_type: 格式类型
            title: 文档标题
            
        Returns:
            格式化器实例
        """
        formatter_class = cls._formatters.get(format_type.lower())
        if not formatter_class:
            raise ValueError(f"不支持的格式: {format_type}")
        
        return formatter_class(title, options=options)
    
    @classmethod
    def get_supported_formats(cls) -> List[str]:
        """获取支持的格式列表"""
        return list(cls._formatters.keys())


